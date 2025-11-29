import json
import re
import urllib.request

from docx import Document
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import Inches

from constants import IMAGE_SECTIONS, IMAGES_MAP
from utils import insert_horizontal_rule, add_link


def generate_long_cv(template, data_file_url, generated_filename):
    generated_doc = Document(template)

    with urllib.request.urlopen(data_file_url) as url:
        data = json.load(url)
        print(data)
        page_names = list(data.keys())
        page_names.remove("header")  # remove the header info because it should not be part of the main loop
        pretty_page_names = {page_name: page_name.replace("_", " ").capitalize() for page_name in page_names}
        print(page_names)
        print(pretty_page_names)

        # header
        generated_doc.add_heading(data["header"]["current_name"], level=0)
        generated_doc.add_paragraph(data["header"]["current_position"], style="Subtitle")
        generated_doc.add_heading(data["header"]["current_address"], level=1)
        generated_doc.add_heading(f"{data["header"]["current_office"]} | {data["header"]["current_emailaddress"]} | {data["header"]["current_orcid"]}", level=2)

        # generate sections for each other page
        for page_name in page_names:
            heading = generated_doc.add_heading()
            run_heading = heading.add_run()
            run_heading.add_picture(f"../images/{IMAGE_SECTIONS[page_name]}.png", width=Inches(0.25))
            run_heading.add_text(f" {pretty_page_names[page_name]}")  # keep a space after the section image
            #generated_doc.add_heading(pretty_page_names[page_name], level=1)
            paragraph_horizontal_rule = generated_doc.add_paragraph()
            insert_horizontal_rule(paragraph_horizontal_rule)
            if page_name == "research_interests":
                # no individual sections for the research interest page
                # however, we still need to add the paragraph
                generated_doc.add_paragraph(data[page_name][0]["descriptions"][0])
            else:
                for section in data[page_name]:
                    if section["title"].lower() == "stay tuned":
                        pass
                    else:
                        # specific titles of the form "title | location     date"
                        if page_name in ["academic_positions", "education", "awards"]:
                            section_title = f"{section["title"]} | {section["subtitles"][0][1]}"
                        else:
                            section_title = section["title"]
                        # regular titles
                        if "date" in section:
                            generated_doc.add_heading(f"{section_title}\t{section["date"]}", level=2)  # four curly-braces: the 2 to show + the 2 to escape the first two
                        else:
                            generated_doc.add_heading(section_title, level=2)  # four curly-braces: the 2 to show + the 2 to escape the first two

                        if "subtitles" in section:
                            if page_name in ["academic_positions", "education", "awards"] and len(section["subtitles"]) == 1 and section["subtitles"][0][0] == "map-pin":
                                # there is only one subtitle and this subtitle is the location
                                # we previously added it to the section title
                                # so we do not create a paragraph for that subtitle
                                # rather, we create paragraphs when there are more subtitles
                                pass
                            else:
                                i = 1
                                paragraph = generated_doc.add_paragraph()
                                for one_subtitle in section["subtitles"]:
                                    if i == 1 and one_subtitle[0] == "map-pin":
                                        pass
                                    else:
                                        subtitle_img = one_subtitle[0]
                                        subtitle_text = one_subtitle[1]
                                        run_text = paragraph.add_run()
                                        # run.add_picture(f"../images/{subtitle_img}.png", width=Inches(0.15))
                                        run_text.add_text(f"{IMAGES_MAP[subtitle_img]}: ")  # Role, Grant, Website, etc
                                        run_text.italic = True
                                        if subtitle_img in ["website", "code-branch"] and subtitle_text.startswith("https://"):
                                            add_link(paragraph, subtitle_text, subtitle_text.replace("https://", "").replace("http://", ""))
                                            #run = paragraph.add_run('www.example.com')
                                            #run.hyperlink.address = 'https://www.example.com'

                                            # _ = add_hyperlink_2(paragraph, subtitle_text, subtitle_text, True)
                                            # run_text_2 = paragraph.add_run()
                                            # run_text_2.add_text(hl)
                                            # add_hyperlink_into_run(paragraph, run_text_2, subtitle_text)
                                            # run_text_2 = add_hyperlink(paragraph, subtitle_text.replace("https://", ""), subtitle_text, False)
                                            # run_text_2 = add_hyperlink(paragraph, subtitle_text, subtitle_text, False)
                                            # run_text_2 = paragraph.add_run()
                                            # run_text_2.add_text(subtitle_text)
                                        else:
                                            run_text_2 = paragraph.add_run()
                                            run_text_2.add_text(f"{subtitle_text}")  # leader, grant, repository, etc
                                        if i < len(section["subtitles"]):
                                            run_text_2.add_text(" | ")
                                    i += 1
                        if "descriptions" in section:
                            if page_name == "publications":
                                # special case: we need to format each publication authors, title, etc
                                for one_description in section["descriptions"]:
                                    format_publication(generated_doc, one_description)
                            elif page_name == "talks":
                                # special case: we need to format each talk
                                for one_description in section["descriptions"]:
                                    format_talk(generated_doc, one_description)
                            else:
                                for one_description in section["descriptions"]:
                                    generated_doc.add_paragraph(f"{one_description}", style='List Bullet')
        generate_files(generated_doc, generated_filename)


def generate_short_cv(template, data_file_url, generated_filename):
    generated_doc = Document(template)

    with urllib.request.urlopen(data_file_url) as url:
        data = json.load(url)
        print(data)
        page_names = list(data.keys())
        page_names.remove("header")  # remove the header info because it should not be part of the main loop
        pretty_page_names = {page_name: page_name.replace("_", " ").capitalize() for page_name in page_names}
        print(page_names)
        print(pretty_page_names)

        # header
        generated_doc.add_heading(data["header"]["current_name"], level=0)
        generated_doc.add_paragraph(data["header"]["current_position"], style="Subtitle")
        generated_doc.add_heading(data["header"]["current_address"], level=1)
        generated_doc.add_heading(f"{data["header"]["current_office"]} | {data["header"]["current_emailaddress"]} | {data["header"]["current_orcid"]}", level=2)

        # generate sections for each other page
        for page_name in page_names:
            heading = generated_doc.add_heading()
            run_heading = heading.add_run()
            run_heading.add_picture(f"../images/{IMAGE_SECTIONS[page_name]}.png", width=Inches(0.25))
            run_heading.add_text(f" {pretty_page_names[page_name]}")  # keep a space after the section image
            #generated_doc.add_heading(pretty_page_names[page_name], level=1)
            paragraph_horizontal_rule = generated_doc.add_paragraph()
            insert_horizontal_rule(paragraph_horizontal_rule)

            if page_name == "research_interests":
                # no individual sections for the research interest page
                # however, we still need to add the paragraph
                generated_doc.add_paragraph(data[page_name][0]["descriptions"][0])
            elif page_name in ["academic_positions", "education", "institutional_responsibilities"]:
                for section in data[page_name]:
                    # specific titles of the form "title | location     date"
                    section_title = f"{section["title"]} | {section["subtitles"][0][1]}"
                    title_as_paragraph = generated_doc.add_paragraph()
                    the_date = " - ".join(re.findall("[0-9]{4}|now", section["date"]))
                    run_date = title_as_paragraph.add_run()
                    run_date.add_text(f"{the_date}\t")
                    run_title = title_as_paragraph.add_run()
                    run_title.add_text(f"{section_title}")
                    run_title.bold = True

                    # we do not print their descriptions
            elif page_name in ["research_visits", "awards", "research_projects", "research_software_and_tools", "working_groups"]:
                i = 0
                title_as_paragraph = generated_doc.add_paragraph()
                for section in data[page_name]:
                    # specific paragraph of the form "title (location. date)"
                    if section["title"].lower() == "stay tuned":
                        pass
                    else:
                        url = None
                        if "subtitles" in section:
                            for subtitle in section["subtitles"]:
                                if subtitle[0] == "code-branch":
                                    url = subtitle[1]
                        if "date" in section:
                            section_title = f"{section["title"]} ({section["date"]})"
                        else:
                            section_title = f"{section["title"]}"
                        if i == len(data[page_name])-1:
                            # last element
                            if url is not None:
                                add_link(title_as_paragraph, url, section_title)
                            else:
                                run_para = title_as_paragraph.add_run()
                                run_para.add_text(f"{section_title}")
                        else:
                            if url is not None:
                                add_link(title_as_paragraph, url, section_title)
                                run_para = title_as_paragraph.add_run()
                                run_para.add_text(f", ")
                            else:
                                run_para = title_as_paragraph.add_run()
                                run_para.add_text(f"{section_title}, ")
                        i += 1

                        # no descriptions here
            elif page_name == "publications":
                # do not print the list of publi but rather the link to my orcid
                publis_para = generated_doc.add_paragraph()
                run_short_publis = publis_para.add_run()
                run_short_publis.add_text(f"My complete list of publications is available on my ")
                add_link(publis_para, f"{data["header"]["current_orcid"]}", "ORCID record.")
            elif page_name in ["professional_service", "reviewing_activities", "teaching_responsibilities", "advising"]:
                for section in data[page_name]:
                    # specific titles of the form "title | location     date"
                    section_title = section["title"]
                    title_as_paragraph = generated_doc.add_paragraph()
                    run_title = title_as_paragraph.add_run()
                    run_title.add_text(f"{section_title}: ")
                    run_title.bold = True
                    i = 1
                    for description in section["descriptions"]:
                        run_descr = title_as_paragraph.add_run()
                        # remove the date before the activity and aggregate all of them
                        if i < len(section["descriptions"]):
                            if page_name in ["teaching_responsibilities", "advising"]:
                                run_descr.add_text(f"{description}, ")
                            else:
                                run_descr.add_text(f"{description[6: len(description)]}, ")
                        else:
                            if page_name in ["teaching_responsibilities", "advising"]:
                                run_descr.add_text(f"{description}")
                            else:
                                run_descr.add_text(f"{description[6: len(description)]}")
            elif page_name in ["talks"]:
                # do not print the list of publi but rather the link to my orcid
                talks_para = generated_doc.add_paragraph()
                run_short_talks = talks_para.add_run()
                run_short_talks.add_text(f"My complete list of talks on my ")
                zenodo_link = "https://zenodo.org/search?q=metadata.creators.person_or_org.name%3A%22Barret%2C%20Nelly%22&l=list&p=1&s=10&sort=bestmatch"
                add_link(talks_para, f"{zenodo_link}", "ZENODO record.")
            # elif page_name in ["teaching_responsibilities", "advising"]:
            #     for section in data[page_name]:
            #         # specific titles of the form "title | location     date"
            #         section_title = section["title"]
            #         title_as_paragraph = generated_doc.add_paragraph()
            #         run_title = title_as_paragraph.add_run()
            #         run_title.add_text(f"{section_title}: ")
            #         for description in section["descriptions"]:
            #             # remove the date
            #             run_title.add_text(description[6: len(description)])
            #             the_date = section["date"]
            #             dates = re.findall("[0-9]{4}|now", the_date)
            #             the_date = " - ".join(dates)
            #
            #             run_date = title_as_paragraph.add_run()
            #             run_date.add_text(f"{the_date}\t")
            #             run_title = title_as_paragraph.add_run()
            #             run_title.add_text(f"{section_title}")
            #             run_title.bold = True
            #         else:
            #             run = title_as_paragraph.add_run()
            #             run.add_text(f"{section_title}")
            #
            #         if "descriptions" in section:
            #             if page_name == "publications":
            #                 # do not print the list of publi but rather the link to my orcid
            #                 publis_para = generated_doc.add_paragraph()
            #                 run_short_publis = publis_para.add_run()
            #                 run_short_publis.add_text(f"Checkout my complete list of publications on my ORCID record: ")
            #                 add_link(publis_para, f"{data["header"]["current_orcid"]}", data["header"]["current_orcid"])
            #             else:
            #                 # do not print descriptions in the short CV
            #                 pass
        generate_files(generated_doc, generated_filename)


def generate_files(document_object, generated_filename):

    docx_filename = f"{generated_filename}.docx"
    pdf_filename = f"{generated_filename}.pdf"
    odt_filename = f"{generated_filename}.odt"

    # Save generated CV into DOCX
    document_object.save(docx_filename)
    if "short" in generated_filename:
        print("Generate short cv: done.")
    else:
        print("Generate long cv: done.")

    # Convert DOCX to ODT
    # converter = Converter(docx_filename)
    # convert_options = WordProcessingConvertOptions()
    # convert_options.format = WordProcessingFileType.ODT
    # converter.convert(odt_filename, convert_options)
    # print("Convert docx to odt: done")

    # convert ODT to PDF
    # convert(docx_filename, pdf_filename)
    # print("Long cv saved as PDF: done.")


def format_publication(document, publi):
    paragraph_item_publi = document.add_paragraph(style='List Number') # numbered list

    # publication authors
    i = 1
    if "main_author" in publi:
        main_author = publi["main_author"]
    else:
        main_author = False
    for author in publi["authors"]:
        run_one_author = paragraph_item_publi.add_run()
        first, last = author["first"], author["last"]
        run_one_author.add_text(f"{first} {last}")
        run_one_author_comma = paragraph_item_publi.add_run()
        # author is a JSON with two fields: first and last names
        if i == len(publi["authors"]):
            # last author
            run_one_author_comma.add_text(f". ")
            if main_author:
                run_one_author.underline = True
                main_author = False
        else:
            run_one_author_comma.add_text(", ")
            if main_author:
                run_one_author.underline = True
                main_author = False
        i += 1

    # publication title
    if "url" in publi:
        add_link(paragraph_item_publi, publi["url"], f"{publi["title"]}")
        run_title = paragraph_item_publi.add_run()
        run_title.add_text(f". ")
    else:
        run_title = paragraph_item_publi.add_run()
        run_title.add_text(f"{publi["title"]}. ")
        run_title.bold = True

    # publication venue
    run_venue = paragraph_item_publi.add_run()
    run_venue.add_text(f"{publi["venue"]}. ")
    run_venue.italic = True

    # publication year
    run_year = paragraph_item_publi.add_run()
    run_year.add_text(f"{publi["year"]}.")
    return paragraph_item_publi


def format_talk(document, talk):
    # restart_numbering(document)
    paragraph_item_talk = document.add_paragraph(style='List Bullet')  # numbered list

    # talk title
    if "url" in talk:
        add_link(paragraph_item_talk, talk["url"], talk["title"])
        run_title = paragraph_item_talk.add_run()
        run_title.add_text(f". ")
    else:
        run_title = paragraph_item_talk.add_run()
        run_title.add_text(f"{talk["title"]}. ")
        run_title.bold = True

    # publication venue
    run_venue = paragraph_item_talk.add_run()
    run_venue.add_text(f"{talk["venue"]}. ")
    run_venue.italic = True

    # publication year
    run_year = paragraph_item_talk.add_run()
    run_year.add_text(f"{talk["year"]}.")
    return paragraph_item_talk


if __name__ == "__main__":
    generate_long_cv("empty-doc-with-styles.docx", "https://nelly-barret.github.io/data/data.json", "cv-long-nelly-barret")
    generate_short_cv("empty-doc-with-styles.docx", "https://nelly-barret.github.io/data/data.json", "cv-short-nelly-barret")