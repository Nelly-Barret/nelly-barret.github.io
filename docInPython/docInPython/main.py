import json

from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches
from urllib.request import urlopen
from bs4 import BeautifulSoup

url = "https://nelly-barret.github.io/RES-projects.html"
html = urlopen(url).read()
soup = BeautifulSoup(html, features="html.parser")

# document = Document()

from docxtpl import DocxTemplate,InlineImage
from docx.shared import Mm

# 1. create template for NINJA, i.e., the Word document wit the NINJA tags ({{tag}})

# 2 fill the template
doc = DocxTemplate('template-jinja.docx')

with open("data.json", "r") as fp:
    data = json.load(fp)
    context = {}
    for header_element in data["header"]:
        context[header_element] = data["header"][header_element]
    i = 0
    for academic_element in data["academic_positions"]:
        context[f"ac_pos{i}_title"] = academic_element["title"]
        context[f"ac_pos{i}_date"] = academic_element["date"]
        j = 0
        for one_description in academic_element["description"]:
            context[f"ac_pos{i}_descr{j}"] = one_description
            j = j+1
        i = i+1
    # context = {'position' : 'INSA LYON',
    #            'address': '7 avenue Jean Capelle',
    #            'pos1_name' : 'Assistant professor',
    #            'pos1_date' : 'Sept. 2025 - now',
    #            'pos1_descr': "bla bla"}
    doc.render(context)
    doc.save("generated-cv.docx")




# doc = DocxTemplate("my_word_template.docx")
# context = { 'company_name' : "World company" }
# doc.render(context)
# doc.save("generated_doc.docx")

# for line in str(soup).split("\n"):
#     stripped_line = line.strip()
#     print(stripped_line)
#     if stripped_line.startswith("<h4"):
#         document.add_heading(stripped_line, level=1)
#     elif stripped_line.startswith("<ul>"):
#         pass
#     elif stripped_line.startswith("<li>"):
#         document.add_paragraph(stripped_line, style='List Bullet')
#
# document.save('demo.docx')

# # kill all script and style elements
# for script in soup(["script", "style"]):
#     script.extract()    # rip it out

# document.add_heading('Document Title', 0)
#
# p = document.add_paragraph('A plain paragraph having some ')
# p.add_run('bold').bold = True
# p.add_run(' and some ')
# p.add_run('italic.').italic = True
#
# document.add_heading('Heading, level 1', level=1)
# document.add_paragraph('Intense quote', style='Intense Quote')
#
# document.add_paragraph(
#     'first item in unordered list', style='List Bullet'
# )
# document.add_paragraph(
#     'first item in ordered list', style='List Number'
# )
#
# #document.add_picture('monty-truth.png', width=Inches(1.25))
#
# records = (
#     (3, '101', 'Spam'),
#     (7, '422', 'Eggs'),
#     (4, '631', 'Spam, spam, eggs, and spam')
# )
#
# table = document.add_table(rows=1, cols=3)
# hdr_cells = table.rows[0].cells
# hdr_cells[0].text = 'Qty'
# hdr_cells[1].text = 'Id'
# hdr_cells[2].text = 'Desc'
# for qty, id, desc in records:
#     row_cells = table.add_row().cells
#     row_cells[0].text = str(qty)
#     row_cells[1].text = id
#     row_cells[2].text = desc
#
# document.add_page_break()
#
# document.save('demo.docx')